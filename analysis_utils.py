from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from mlxtend.frequent_patterns import apriori, association_rules
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import pandas as pd


def basic_stats(df, save_path=None):
    df['InvoiceDate'] = pd.to_datetime(df['InvoiceDate'])
    df['TotalValue'] = df['Quantity'] * df['Price']

    if save_path:
        os.makedirs(save_path, exist_ok=True)

    grouped = df.groupby("Invoice")["TotalValue"].sum()
    profit = df.groupby("Description")["TotalValue"].sum().sort_values(ascending=False)
    top_quantity = df.groupby("Description")["Quantity"].sum().sort_values(ascending=False)

    daily = df.groupby(df['InvoiceDate'].dt.date).agg({
        'Invoice': 'nunique',
        'Quantity': 'sum',
        'TotalValue': 'sum'
    }).rename(columns={
        'Invoice': 'Liczba transakcji',
        'Quantity': 'Łączna ilość',
        'TotalValue': 'Przychód dzienny'
    })

    print("=== Podstawowe statystyki wyświetlone w terminalu ===")
    print("Okres:", df["InvoiceDate"].min().date(), "→", df["InvoiceDate"].max().date())
    print("Unikalni klienci:", df["Customer ID"].nunique())
    print("Unikalne produkty:", df["Description"].nunique())
    print("Liczba krajów:", df["Country"].nunique())
    print("\nTop 5 produkty wg ilości:\n", top_quantity.head(5))
    print("\nTop 5 produkty wg wartości:\n", profit.head(5))
    print("\nPrzykładowa dzienna agregacja:\n", daily.head(5))

    if save_path:
        stats_summary = {
            "Okres od": [df["InvoiceDate"].min().date()],
            "Okres do": [df["InvoiceDate"].max().date()],
            "Liczba transakcji": [df["Invoice"].nunique()],
            "Liczba klientów": [df["Customer ID"].nunique()],
            "Liczba produktów": [df["Description"].nunique()],
            "Liczba krajów": [df["Country"].nunique()],
            "Średnia wartość transakcji": [round(grouped.mean(), 2)],
            "Mediana transakcji": [round(grouped.median(), 2)],
            "Maksymalna wartość transakcji": [round(grouped.max(), 2)],
            "Minimalna wartość transakcji": [round(grouped.min(), 2)]
        }
        summary_df = pd.DataFrame(stats_summary)
        summary_df.to_csv(f"{save_path}/podsumowanie_statystyk.csv", index=False)
        top_quantity.head(10).to_csv(f"{save_path}/top_produkty_ilosc.csv")
        profit.head(10).to_csv(f"{save_path}/top_produkty_zysk.csv")
        daily.to_csv(f"{save_path}/dzienna_zmiennosc.csv")

        print(f"\nStatystyki zapisane do katalogu: {save_path}")

    return daily

def get_top_revenue_countries(df, top_n=12):
    df["TotalValue"] = df["Quantity"] * df["Price"]
    top_countries = (
        df.groupby("Country")["TotalValue"]
        .sum()
        .sort_values(ascending=False)
        .head(top_n)
        .index.tolist()
    )
    return top_countries

def prepare_basket(df, allowed_items):
    df = df[df["Description"].isin(allowed_items)]
    basket = df.groupby(['Invoice', 'Description'])['Quantity'].sum().unstack().fillna(0)
    basket = basket.gt(0)
    return basket

def run_apriori_analysis(df, country_list, output_dir="outputs_koszykowa"):
    os.makedirs(output_dir, exist_ok=True)

    product_counts = df["Description"].value_counts()
    support_threshold = product_counts.quantile(0.80)
    allowed_items = product_counts[product_counts >= support_threshold].index.tolist()

    df_filtered_global = df[df["Description"].isin(allowed_items)]
    global_basket = df_filtered_global.groupby(['Invoice', 'Description'])['Quantity'].sum().unstack().fillna(0)
    global_basket = global_basket.gt(0)

    if global_basket.empty:
        print("Globalny koszyk pusty — nie można wyliczyć min_support.")
        return

    min_sup = 0.2

    print(f"Globalny min_support ustawiony na: {round(min_sup, 4)}")

    for country in country_list:
        print(f"\n=== 🇨🇭 Analiza koszykowa dla: {country} ===")
        df_country = df[df["Country"] == country]
        df_country = df_country[df_country["Description"].isin(allowed_items)]

        basket = df_country.groupby(['Invoice', 'Description'])['Quantity'].sum().unstack().fillna(0)
        basket = basket.gt(0)

        if basket.empty or basket.shape[1] < 2:
            print(f"Pomijam {country} – za mało danych po filtracji.")
            continue

        try:
            freq_items = apriori(basket, min_support=min_sup, use_colnames=True)
        except Exception as e:
            print(f"Błąd w apriori() dla {country}: {e}")
            continue

        if freq_items.empty:
            print(f"Brak zbiorów częstych dla {country}.")
            continue

        try:
            rules = association_rules(freq_items, metric="confidence", min_threshold=0.8)
        except Exception as e:
            print(f"Błąd w association_rules() dla {country}: {e}")
            continue

        freq_path = os.path.join(output_dir, f"{country}_frequent_itemsets.csv")
        rules_path = os.path.join(output_dir, f"{country}_association_rules.csv")
        freq_items.to_csv(freq_path, index=False)
        rules.to_csv(rules_path, index=False)

        print(f"{country}: {len(freq_items)} zbiorów częstych, {len(rules)} reguł")



def plot_all_basic_stats(df, daily, save_path="outputs"):
    os.makedirs(save_path, exist_ok=True)
    sns.set_style("whitegrid")

    plt.figure(figsize=(12, 5))
    sns.lineplot(x=daily.index, y=daily["Liczba transakcji"], marker="o", linewidth=2)
    plt.title("Liczba transakcji dziennie")
    plt.xlabel("Data")
    plt.ylabel("Transakcje")
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.savefig(f"{save_path}/transakcje_dziennie.png")
    plt.close()

    plt.figure(figsize=(12, 5))
    sns.lineplot(x=daily.index, y=daily["Przychód dzienny"], marker="x", color="green", linewidth=2)
    plt.title("Przychód dzienny")
    plt.xlabel("Data")
    plt.ylabel("Przychód")
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.savefig(f"{save_path}/przychod_dzienny.png")
    plt.close()

    top_qty = df.groupby("Description")["Quantity"].sum().sort_values(ascending=False).head(10)
    plt.figure(figsize=(10, 6))
    sns.barplot(x=top_qty.values, y=top_qty.index, palette="Blues_d")
    plt.title("Top 10 produktów wg ilości")
    plt.xlabel("Suma ilości")
    plt.tight_layout()
    plt.savefig(f"{save_path}/top_produkty_ilosc.png")
    plt.close()

    df["TotalValue"] = df["Quantity"] * df["Price"]
    top_val = df.groupby("Description")["TotalValue"].sum().sort_values(ascending=False).head(10)
    plt.figure(figsize=(10, 6))
    sns.barplot(x=top_val.values, y=top_val.index, palette="Oranges_r")
    plt.title("Top 10 produktów wg przychodu")
    plt.xlabel("Suma przychodu")
    plt.tight_layout()
    plt.savefig(f"{save_path}/top_produkty_przychod.png")
    plt.close()

    top_countries = df.groupby("Country")["TotalValue"].sum().sort_values(ascending=False).head(10)
    plt.figure(figsize=(10, 6))
    sns.barplot(x=top_countries.values, y=top_countries.index, palette="viridis")
    plt.title("Top 10 krajów wg przychodu")
    plt.xlabel("Łączny przychód")
    plt.tight_layout()
    plt.savefig(f"{save_path}/top_kraje_przychod.png")
    plt.close()

    transactions_by_country = df.groupby("Country")["Invoice"].nunique().sort_values(ascending=False).head(10)
    plt.figure(figsize=(10, 6))
    sns.barplot(x=transactions_by_country.values, y=transactions_by_country.index, palette="mako")
    plt.title("Liczba transakcji na kraj (Top 10)")
    plt.xlabel("Liczba transakcji")
    plt.tight_layout()
    plt.savefig(f"{save_path}/top_kraje_transakcje.png")
    plt.close()

    print(f"Wykresy zapisane do katalogu: {save_path}")

def generate_word_report(df, daily, outputs_dir="WDED_SOLO", koszykowa_dir="outputs/Market Basket Analysis"):
    os.makedirs(outputs_dir, exist_ok=True)

    start = df["InvoiceDate"].min().date()
    end = df["InvoiceDate"].max().date()
    total_rows = len(df)
    unique_customers = df["Customer ID"].nunique()
    unique_products = df["Description"].nunique()
    unique_countries = df["Country"].nunique()
    num_invoices = df["Invoice"].nunique()
    mean_trans = round(df.groupby("Invoice")["TotalValue"].sum().mean(), 2)

    top_qty = df.groupby("Description")["Quantity"].sum().sort_values(ascending=False).head(10)
    top_val = df.groupby("Description")["TotalValue"].sum().sort_values(ascending=False).head(10)
    top_countries = df.groupby("Country")["TotalValue"].sum().sort_values(ascending=False).head(10)
    country_tx = df.groupby("Country")["Invoice"].nunique().sort_values(ascending=False).head(10)

    min_day = daily["Liczba transakcji"].min()
    max_day = daily["Liczba transakcji"].max()
    mean_day = int(daily["Liczba transakcji"].mean())
    min_revenue = int(daily["Przychód dzienny"].min())
    max_revenue = int(daily["Przychód dzienny"].max())
    mean_revenue = int(daily["Przychód dzienny"].mean())

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    doc.add_heading('Raport z analizy danych zakupów online', 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_heading('1. Wprowadzenie', level=1)
    doc.add_paragraph(
        "W dobie rosnącego znaczenia e-commerce analiza danych zakupowych staje się kluczowym elementem zrozumienia potrzeb klientów "
        "i optymalizacji procesów sprzedażowych. Niniejszy raport stanowi wynik eksploracyjnej analizy danych transakcyjnych pochodzących "
        "ze sklepu internetowego, zrealizowanych w latach 2009–2010. Analiza obejmowała czyszczenie danych, wyznaczenie kluczowych statystyk, "
        "a także wykorzystanie algorytmu Apriori do wykrywania reguł asocjacyjnych między produktami."
    )

    doc.add_heading('2. Charakterystyka danych', level=1)
    doc.add_paragraph(
        f"Zakres czasowy danych: {start} - {end}. Rekordy: {total_rows}, Transakcje: {num_invoices}, Klienci: {unique_customers}, Produkty: {unique_products}, Kraje: {unique_countries}, \n"
        f"Średnia wartość transakcji: {mean_trans} GBP."
    )

    doc.add_heading('3. Najpopularniejsze produkty i rynki', level=1)

    doc.add_paragraph("Top 10 produktów wg ilości:")
    table1 = doc.add_table(rows=1, cols=2)
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Produkt'
    hdr_cells[1].text = 'Suma ilości'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '2F75B5')  # niebieski
        cell._tc.get_or_add_tcPr().append(shd)

    for desc, qty in top_qty.items():
        row_cells = table1.add_row().cells
        row_cells[0].text = str(desc)
        row_cells[1].text = str(int(qty))

    doc.add_paragraph()

    doc.add_paragraph("Top 10 produktów wg przychodu:")
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = 'Table Grid'
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = 'Produkt'
    hdr_cells[1].text = 'Przychód [GBP]'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'C00000')  # czerwony
        cell._tc.get_or_add_tcPr().append(shd)

    for desc, val in top_val.items():
        row_cells = table2.add_row().cells
        row_cells[0].text = str(desc)
        row_cells[1].text = f"{val:,.2f}"

    doc.add_paragraph()

    doc.add_paragraph("Top 10 krajów wg przychodu:")
    table3 = doc.add_table(rows=1, cols=2)
    table3.style = 'Table Grid'
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = 'Kraj'
    hdr_cells[1].text = 'Przychód [GBP]'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '548235')  # zielony
        cell._tc.get_or_add_tcPr().append(shd)

    for country, val in top_countries.items():
        row_cells = table3.add_row().cells
        row_cells[0].text = str(country)
        row_cells[1].text = f"{val:,.2f}"

    doc.add_paragraph()

    doc.add_paragraph("Top 10 krajów wg liczby transakcji:")
    table4 = doc.add_table(rows=1, cols=2)
    table4.style = 'Table Grid'
    hdr_cells = table4.rows[0].cells
    hdr_cells[0].text = 'Kraj'
    hdr_cells[1].text = 'Liczba transakcji'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'F79646')  # pomarańczowy
        cell._tc.get_or_add_tcPr().append(shd)

    for country, tx in country_tx.items():
        row_cells = table4.add_row().cells
        row_cells[0].text = str(country)
        row_cells[1].text = str(tx)

    doc.add_paragraph()

    doc.add_heading('4. WykresyD', level=1)

    for name in ["transakcje_dziennie", "przychod_dzienny", "top_produkty_ilosc", "top_produkty_przychod", "top_kraje_przychod", "top_kraje_transakcje"]:
        image_path = os.path.join(outputs_dir, f"{name}.png")
        if os.path.exists(image_path):
            doc.add_picture(image_path, width=Inches(6))

    doc.add_heading('5. Reguły asocjacyjne (analiza koszykowa)', level=1)
    koszyk_info = ""
    if os.path.isdir(koszykowa_dir):
        files = sorted([f for f in os.listdir(koszykowa_dir) if f.endswith("_frequent_itemsets.csv")])
        for f in files:
            country = f.split("_frequent_itemsets.csv")[0]
            rules_file = os.path.join(koszykowa_dir, f"{country}_association_rules.csv")
            rules_count = 0
            if os.path.exists(rules_file):
                rules_df = pd.read_csv(rules_file)
                rules_count = len(rules_df)
            freq_df = pd.read_csv(os.path.join(koszykowa_dir, f))
            koszyk_info += f"{country}: {len(freq_df)} zbiorów częstych, {rules_count} reguł\n"
    doc.add_paragraph(koszyk_info.strip() if koszyk_info else "Brak danych koszykowych.")

    doc.add_heading('6. Wykorzystane narzędzia i podejście techniczne', level=1)
    doc.add_paragraph(
        "Analizę wykonano w języku Python z użyciem bibliotek: pandas, matplotlib, seaborn, mlxtend, openpyxl oraz python-docx."
        " Dane wczytano z Excela, oczyszczono i poddano analizie koszykowej metodą Apriori. Raport i wykresy wygenerowano automatycznie."
    )

    doc.add_heading('7. Wnioski i rekomendacje', level=1)
    doc.add_paragraph(
        "Przeprowadzona analiza ujawniła wyraźną dominację Zjednoczonego Królestwa jako kluczowego rynku sprzedażowego — zarówno pod względem liczby transakcji, jak i generowanego przychodu. "
        "Zdecydowana większość zamówień pochodziła właśnie z tego kraju, co może wskazywać na silnie ugruntowaną bazę klientów oraz skuteczne działania marketingowe prowadzone lokalnie.\n"
    )
    doc.add_paragraph(
        "Wśród produktów największą popularnością cieszyły się przedmioty dekoracyjne, takie jak świeczniki czy ozdobne akcesoria kuchenne. Świadczy to o popycie na estetykę oraz personalizację przestrzeni domowej wśród klientów e-commerce. "
        "Z kolei w analizie przychodów, najwyższe wartości osiągały również produkty o charakterze użytkowym, lecz premium, co sugeruje istnienie segmentu klientów gotowych do ponoszenia większych wydatków za wyższą jakość lub estetykę."
    )
    doc.add_paragraph(
        "Oczyszczenie danych z błędnych lub niepełnych rekordów (np. zwrotów i niezapłaconych zamówień) okazało się kluczowe dla jakości analizy i poprawności wygenerowanych modeli. "
        "Bez tego etapu istniałoby ryzyko błędnych wniosków, zwłaszcza przy analizie reguł asocjacyjnych.\n"
    )
    doc.add_paragraph(
        "Analiza koszykowa przeprowadzona na 12 krajach o najwyższym przychodzie wykazała istnienie wielu zbiorów częstych, z których część prowadziła do znaczących reguł asocjacyjnych. "
        "Choć nie wszystkie kraje wykazały wystarczające powiązania między produktami, uzyskane wyniki mogą zostać z powodzeniem wykorzystane w implementacji mechanizmów rekomendacyjnych (np. systemów typu \"często kupowane razem\"), "
        "a także w działaniach cross-sellingowych i tworzeniu spersonalizowanych ofert produktowych."
    )
    doc.add_paragraph(
        "Zebrane dane oraz przeprowadzone analizy mogą stanowić solidną podstawę do wdrożenia bardziej zaawansowanych technik predykcyjnych, takich jak segmentacja klientów, prognozowanie popytu, czy też dynamiczne ustalanie cen. "
        "Rekomenduje się kontynuację badań w kierunku analizy sezonowości zakupów, a także integracji danych z kanałów marketingowych w celu pełniejszego zrozumienia ścieżki zakupowej klienta."
    )

    path = os.path.join(outputs_dir, "raport.docx")
    doc.save(path)
    print(f"\u2705 Raport zapisany do: {path}")
    return path






